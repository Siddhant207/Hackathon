import io
import pypdf
import docx

def parse_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF file bytes."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    extracted_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            extracted_text.append(text)
    return "\n".join(extracted_text).strip()

def parse_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX file bytes."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs).strip()

def parse_txt(file_bytes: bytes) -> str:
    """Extract text from TXT bytes."""
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore").strip()

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Detect file extension and extract plain text."""
    lower_filename = filename.lower()
    if lower_filename.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif lower_filename.endswith(".docx") or lower_filename.endswith(".doc"):
        return parse_docx(file_bytes)
    elif lower_filename.endswith(".txt"):
        return parse_txt(file_bytes)
    else:
        # Default try utf-8
        return parse_txt(file_bytes)
